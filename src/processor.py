from docx import Document
from docx2pdf import convert
from datetime import datetime
import os
import sys
from reportlab.lib.pagesizes import A4
from reportlab.platypus import Table, TableStyle, Paragraph, SimpleDocTemplate, Spacer
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.units import inch, mm

def get_template_path(template_path):
    """获取模板文件的绝对路径，兼容开发和打包环境"""
    if getattr(sys, 'frozen', False):
        base_path = os.path.dirname(sys.executable)
        return os.path.join(base_path, 'templates', os.path.basename(template_path))
    else:
        current_dir = os.path.dirname(os.path.abspath(__file__))
        return os.path.join(os.path.dirname(current_dir), 'templates', os.path.basename(template_path))

def fill_word_template(template_path, data):
    """填充 Word 模板中的占位符并保存"""
    absolute_template_path = get_template_path(template_path)
    print(f"使用模板文件路径: {absolute_template_path}")
    print(f"替换数据: {data}")
    
    if not os.path.exists(absolute_template_path):
        raise FileNotFoundError(f"找不到模板文件: {absolute_template_path}")
    
    doc = Document(absolute_template_path)
    
    for paragraph in doc.paragraphs:
        replace_placeholder_in_paragraph(paragraph, data)
    
    date_str = datetime.now().strftime("%Y%m%d")
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    base_name = os.path.basename(template_path).replace(".docx", "").replace("_Template", "")
    output_dir = os.path.join("output", date_str)
    os.makedirs(output_dir, exist_ok=True)
    seller_name = data.get('seller_name', '').replace('/', '_').replace('\\', '_').replace(':', '_')[:50]
    file_prefix = f"{seller_name}-{base_name}_{timestamp}" if seller_name else f"{base_name}_{timestamp}"
    output_path = os.path.join(output_dir, f"{file_prefix}_filled.docx")
    doc.save(output_path)
    return output_path

def convert_to_pdf(doc_path, output_pdf_path):
    """将 Word 文件转换为 PDF"""
    convert(doc_path, output_pdf_path)

def generate_invoice_pdf(data, output_path):
    """生成发票 PDF 文件"""
    pdfmetrics.registerFont(TTFont('SimSun', 'C:\\Windows\\Fonts\\simsun.ttc'))
    pdfmetrics.registerFont(TTFont('Microsoft YaHei', 'C:\\Windows\\Fonts\\msyh.ttc'))

    doc = SimpleDocTemplate(output_path, pagesize=A4, leftMargin=15*mm, rightMargin=15*mm, topMargin=15*mm, bottomMargin=15*mm)
    elements = []
    styles = getSampleStyleSheet()

    title_style = ParagraphStyle('Title', parent=styles['Normal'], fontName='Microsoft YaHei', fontSize=36, alignment=1, textColor=colors.navy, spaceAfter=10, spaceBefore=10, leading=32, fontWeight='bold')
    address_style = ParagraphStyle('Address', parent=styles['Normal'], fontName='Microsoft YaHei', fontSize=10, leading=12, leftIndent=3, rightIndent=3, wordWrap='CJK', alignment=0, textColor=colors.navy, fontWeight='bold')
    normal_style = ParagraphStyle('Normal', parent=styles['Normal'], fontName='Microsoft YaHei', fontSize=9, leading=11, leftIndent=3, rightIndent=3, textColor=colors.navy, fontWeight='bold')

    elements.append(Paragraph("Invoice", title_style))
    elements.append(Spacer(1, 0.1 * inch))

    address_and_info_data = [
        [
            [
                Paragraph("From:", normal_style),
                Paragraph(data['invoice_my_address'].replace(', ', ',<br/>'), address_style),
                Spacer(1, 0.08 * inch),
                Paragraph("To:", normal_style),
                Paragraph(data['invoice_address'].replace(', ', ',<br/>'), address_style)
            ],
            [
                Table(
                    [["Invoice No:", data['invoice_no']], ["Invoice Date:", format_date_to_english(data['invoice_date'])]],
                    colWidths=[0.8 * inch, 1.8 * inch],
                    style=TableStyle([
                        ('ALIGN', (0, 0), (0, -1), 'LEFT'),
                        ('ALIGN', (1, 0), (1, -1), 'LEFT'),
                        ('FONTNAME', (0, 0), (-1, -1), 'Microsoft YaHei'),
                        ('FONTSIZE', (0, 0), (-1, -1), 9),
                        ('LEFTPADDING', (0, 0), (-1, -1), 3),
                        ('RIGHTPADDING', (0, 0), (-1, -1), 3),
                        ('TOPPADDING', (0, 0), (-1, -1), 1),
                        ('BOTTOMPADDING', (0, 0), (-1, -1), 1),
                        ('TEXTCOLOR', (0, 0), (-1, -1), colors.navy),
                        ('FONTWEIGHT', (0, 0), (-1, -1), 'bold')
                    ])
                )
            ]
        ]
    ]
    
    page_width = A4[0] - 30*mm
    address_table = Table(address_and_info_data, colWidths=[0.6 * page_width, 0.4 * page_width])
    address_table.setStyle(TableStyle([('VALIGN', (0, 0), (-1, -1), 'TOP'), ('LEFTPADDING', (0, 0), (-1, -1), 0), ('RIGHTPADDING', (0, 0), (-1, -1), 0), ('TOPPADDING', (0, 0), (-1, -1), 0), ('BOTTOMPADDING', (0, 0), (-1, -1), 0)]))
    elements.append(address_table)
    elements.append(Spacer(1, 0.4 * inch))

    template_type = data.get('template_type', 'Invoice_AD_Template.docx')
    table_data, merge_info = process_table_data(data, template_type)

    if template_type == 'Invoice_Promo_Template.docx':
        table = Table(table_data, colWidths=[page_width*0.40, page_width*0.18, page_width*0.27, page_width*0.15], rowHeights=[0.3 * inch] * len(table_data))
    else:
        table = Table(table_data, colWidths=[page_width*0.35, page_width*0.18, page_width*0.15, page_width*0.15, page_width*0.17], rowHeights=[0.3 * inch] * len(table_data))

    table_style = [
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('ALIGN', (0, 0), (0, -1), 'CENTER'),
        ('ALIGN', (1, 0), (3, -1), 'CENTER'),
        ('ALIGN', (-1, 0), (-1, -1), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('LEFTPADDING', (0, 0), (-1, -1), 5),
        ('RIGHTPADDING', (0, 0), (-1, -1), 5),
        ('TOPPADDING', (0, 0), (-1, -1), 5),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.navy),
        ('FONTNAME', (0, 0), (-1, -1), 'Microsoft YaHei'),
        ('FONTSIZE', (0, 0), (-1, 0), 9),
        ('FONTSIZE', (0, 1), (-1, -1), 9),
        ('BACKGROUND', (0, 0), (-1, 0), colors.Color(0.9, 0.9, 0.95)),
    ]

    if 'table_data' in data and data['table_data'].strip():
        if template_type == 'Invoice_Promo_Template.docx':
            for desc, indices in merge_info['description'].items():
                apply_span(table_style, indices, 0)
            for asin, indices in merge_info['asin'].items():
                apply_span(table_style, indices, 1)
            for product, indices in merge_info['product'].items():
                apply_span(table_style, indices, 2)
            if 'amount' in merge_info:
                for amount, indices in merge_info['amount'].items():
                    apply_span(table_style, indices, 3)
        else:
            for desc, indices in merge_info['description'].items():
                apply_span(table_style, indices, 0)
            for asin, indices in merge_info['asin'].items():
                apply_span(table_style, indices, 1)

    table.setStyle(TableStyle(table_style))
    elements.append(table)
    elements.append(Spacer(1, 0.4 * inch))

    bank_info_table = create_info_table(data['bank_info'], "Bank Info: 美元账户", page_width)
    elements.append(bank_info_table)
    elements.append(Spacer(1, 0.4 * inch))

    remarks_table = create_info_table({"remark": "我司不承担任何手续费"}, "备注:", page_width)
    elements.append(remarks_table)

    doc.build(elements)

def apply_span(table_style, indices, col):
    """应用表格单元格合并"""
    indices.sort()
    start_idx = indices[0]
    for i in range(1, len(indices)):
        if indices[i] != indices[i-1] + 1:
            if start_idx != indices[i-1]:
                table_style.append(('SPAN', (col, start_idx), (col, indices[i-1])))
            start_idx = indices[i]
    if start_idx != indices[-1] and len(indices) > 1:
        table_style.append(('SPAN', (col, start_idx), (col, indices[-1])))

def log_submission(data, output_file):
    """记录提交日志"""
    timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    log_entry = f"[{timestamp}] 生成文件: {output_file}, 数据: {data}\n"
    log_dir = "logs"
    os.makedirs(log_dir, exist_ok=True)
    with open(os.path.join(log_dir, "submission_log.txt"), "a", encoding="utf-8") as f:
        f.write(log_entry)

def format_date_to_english(date_str):
    """将日期字符串格式化为英文格式"""
    if isinstance(date_str, datetime):
        return date_str.strftime('%B %d, %Y')
    date_str = str(date_str)
    date_formats = ['%Y-%m-%d', '%Y/%m/%d', '%d/%m/%Y', '%m/%d/%Y', '%Y.%m.%d']
    for date_format in date_formats:
        try:
            date_obj = datetime.strptime(date_str, date_format)
            return date_obj.strftime('%B %d, %Y')
        except ValueError:
            continue
    return date_str

def replace_placeholder_in_paragraph(paragraph, data):
    """替换段落中的占位符，同时精确控制格式"""
    has_placeholder = False
    for key in data.keys():
        if f"{{{{{key}}}}}" in paragraph.text:
            has_placeholder = True
            break
    if not has_placeholder:
        return

    # 收集原始 runs 的格式信息
    runs_info = []
    for run in paragraph.runs:
        runs_info.append({
            'text': run.text,
            'bold': run.bold,
            'italic': run.italic,
            'underline': run.underline,
            'font_name': run.font.name if run.font.name else None,
            'font_size': run.font.size if hasattr(run.font, 'size') else None,
            'font_color': run.font.color.rgb if hasattr(run.font, 'color') and run.font.color.rgb else None
        })

    # 获取原始完整文本并进行替换，同时记录替换位置
    full_text = paragraph.text
    placeholder_positions = {}  # 记录每个替换值的起始和结束位置
    for key, value in data.items():
        placeholder = f"{{{{{key}}}}}"
        if placeholder in full_text:
            print(f"在段落中找到占位符: {placeholder}, 替换为: {value}")
            start = full_text.find(placeholder)
            end = start + len(placeholder)
            full_text = full_text.replace(placeholder, str(value))
            # 调整位置，记录新值的范围
            new_start = start
            new_end = start + len(str(value))
            placeholder_positions[(new_start, new_end)] = (key, str(value))

    # 清空原有 runs
    for i in range(len(paragraph.runs) - 1, -1, -1):
        paragraph._p.remove(paragraph.runs[i]._r)

    # 重建 runs 并精确应用格式
    current_pos = 0
    while current_pos < len(full_text):
        # 检查当前字符是否在某个替换值的范围内
        in_placeholder = False
        placeholder_key = None
        for (start, end), (key, value) in placeholder_positions.items():
            if current_pos >= start and current_pos < end:
                in_placeholder = True
                placeholder_key = key
                segment_end = min(end, len(full_text))
                break
        if not in_placeholder:
            # 查找下一个占位符的起始位置或文本末尾
            segment_end = len(full_text)
            for (start, _), _ in placeholder_positions.items():
                if start > current_pos:
                    segment_end = min(segment_end, start)
            # 使用原始 runs 的格式
            segment_text = full_text[current_pos:segment_end]
            new_run = paragraph.add_run(segment_text)
            for run_info in runs_info:
                if run_info['text'] in paragraph.text:
                    new_run.bold = run_info['bold']
                    new_run.italic = run_info['italic']
                    new_run.underline = run_info['underline']
                    if run_info['font_name']:
                        new_run.font.name = run_info['font_name']
                    if run_info['font_size']:
                        new_run.font.size = run_info['font_size']
                    if run_info['font_color']:
                        new_run.font.color.rgb = run_info['font_color']
                    break
        else:
            # 处理替换值，精确应用格式
            segment_text = full_text[current_pos:segment_end]
            new_run = paragraph.add_run(segment_text)
            # 根据字段类型决定是否应用下划线
            if placeholder_key != "business_name":  # 示例：business_name 不需要下划线
                new_run.underline = True
            # 其他格式可以根据需要继承或自定义
        current_pos = segment_end

def process_table_data(data, template_type):
    """处理表格数据，返回表格内容和合并信息"""
    table_data = []
    merge_info = {}
    total_amount = 0

    if 'table_data' in data and data['table_data'].strip():
        rows = data['table_data'].replace('\r\n', '\n').strip().split('\n')
        if template_type == 'Invoice_Promo_Template.docx':
            table_data.append(["Description", "Asin", "Product", "Amount"])
            merge_info = {'description': {}, 'asin': {}, 'product': {}, 'amount': {}}
            unique_combinations = set()
            for i, row in enumerate(rows):
                if row.strip():
                    parts = [part.strip() for part in row.split(',')]
                    if len(parts) >= 4:
                        try:
                            description, asin, product, amount = parts[0], parts[1], parts[2], float(parts[3])
                            combo_key = f"{product}|{amount}"
                            if combo_key not in unique_combinations:
                                unique_combinations.add(combo_key)
                                total_amount += amount
                            table_data.append([description, asin, product, f"{amount:.2f}"])
                            row_idx = len(table_data) - 1
                            for key, value in zip(['description', 'asin', 'product', 'amount'], [description, asin, product, f"{amount:.2f}"]):
                                if value not in merge_info[key]:
                                    merge_info[key][value] = []
                                merge_info[key][value].append(row_idx)
                        except (ValueError, IndexError) as e:
                            print(f"处理行数据时出错: {e}")
                            continue
        else:
            table_data.append(["Description", "Asin", "Per day", "Day", "AMOUNT"])
            merge_info = {'description': {}, 'asin': {}}
            for i, row in enumerate(rows):
                if row.strip():
                    parts = [part.strip() for part in row.split(',')]
                    if len(parts) >= 4:
                        try:
                            description, asin, rate, days = parts[0], parts[1], float(parts[2]), int(parts[3])
                            row_amount = rate * days
                            total_amount += row_amount
                            table_data.append([description, asin, f"{rate:.2f}", f"{days} days", f"{row_amount:.2f}"])
                            row_idx = len(table_data) - 1
                            for key, value in zip(['description', 'asin'], [description, asin]):
                                if value not in merge_info[key]:
                                    merge_info[key][value] = []
                                merge_info[key][value].append(row_idx)
                        except (ValueError, IndexError) as e:
                            print(f"处理行数据时出错: {e}")
                            continue
    data['amount'] = str(total_amount)
    if template_type == 'Invoice_Promo_Template.docx':
        table_data.append(["TOTAL", "", "", f"USD {total_amount:.2f}"])
    else:
        table_data.append(["TOTAL", "", "", "", f"USD {total_amount:.2f}"])
    return table_data, merge_info

def create_info_table(info_data, title, page_width):
    """生成信息表格（如银行信息或备注）"""
    data = [[title, ""]]
    if title == "Bank Info: 美元账户":
        data.append(["(接受银行转账,及第三方如PingPong,空中云汇,万里汇等转账.)", ""])
        for key, value in info_data.items():
            data.append([f"{key.replace('_', ' ').title()}: {value}", ""])
    else:
        data.append([info_data['remark'], ""])
    table = Table(data, colWidths=[page_width * 0.25, page_width * 0.75])
    table.setStyle(TableStyle([
        ('ALIGN', (0, 0), (0, -1), 'LEFT'),
        ('ALIGN', (1, 0), (1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (-1, -1), 'Microsoft YaHei'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('LEFTPADDING', (0, 0), (0, -1), 3),
        ('RIGHTPADDING', (0, 0), (0, -1), 10),
        ('LEFTPADDING', (1, 0), (1, -1), 10),
        ('RIGHTPADDING', (1, 0), (1, -1), 3),
        ('TOPPADDING', (0, 0), (-1, -1), 2),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 2),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.navy),
        ('FONTWEIGHT', (0, 0), (-1, -1), 'bold'),
        ('GRID', (0, 0), (-1, -1), 0, colors.white),
    ]))
    return table

if __name__ == "__main__":
    template_path = "templates/example.docx"
    data = {"name": "张三", "date": "2023-10-01"}
    filled_doc_path = fill_word_template(template_path, data)
    pdf_timestamp = os.path.basename(filled_doc_path).split("_")[-2]
    output_pdf_path = filled_doc_path.replace("_filled.docx", f"_{pdf_timestamp}.pdf")
    convert_to_pdf(filled_doc_path, output_pdf_path)
    log_submission(data, output_pdf_path)